#!/usr/bin/env python3
"""
JD信息提取处理器
从多种格式文件中提取JD信息并生成Excel表格
"""

import os
import sys
import json
import glob
from datetime import datetime
from typing import List, Dict, Optional, Tuple
import pandas as pd
from pathlib import Path

# 输出字段定义
OUTPUT_COLUMNS = [
    "岗位分类",
    "公司名称",
    "部门名称",
    "职位名称",
    "薪资范围",
    "工作地点",
    "通用硬性要求",
    "HR对焦信息描述",
    "JD岗位开放日期",
    "来源文件名"
]

# 支持的文件格式
SUPPORTED_EXTENSIONS = {
    'excel': ['.xlsx', '.xls'],
    'word': ['.docx'],
    'image': ['.png', '.jpg', '.jpeg'],
    'text': ['.txt']
}


def get_file_type(filepath: str) -> Optional[str]:
    """根据文件扩展名判断文件类型"""
    ext = Path(filepath).suffix.lower()
    for file_type, extensions in SUPPORTED_EXTENSIONS.items():
        if ext in extensions:
            return file_type
    return None


def scan_directory(directory: str) -> Dict[str, List[str]]:
    """扫描目录下的所有支持格式的文件"""
    files_by_type = {
        'excel': [],
        'word': [],
        'image': [],
        'text': [],
        'unsupported': []
    }

    for filepath in glob.glob(os.path.join(directory, '*')):
        if os.path.isfile(filepath):
            file_type = get_file_type(filepath)
            if file_type:
                files_by_type[file_type].append(filepath)
            else:
                files_by_type['unsupported'].append(filepath)

    return files_by_type


def extract_from_excel(filepath: str) -> Tuple[List[Dict], Optional[str]]:
    """
    从Excel文件提取JD信息
    返回: (职位列表, 错误信息)
    """
    jobs = []
    try:
        xl = pd.ExcelFile(filepath)

        for sheet_name in xl.sheet_names:
            df = pd.read_excel(filepath, sheet_name=sheet_name)

            # 识别包含JD信息的列
            # 常见列名关键词
            jd_keywords = ['职位', '岗位', '要求', '职责', '描述', '地点', '薪资', '公司', '部门']

            # 检查是否有JD相关的列
            has_jd_columns = any(
                any(kw in str(col) for kw in jd_keywords)
                for col in df.columns
            )

            if not has_jd_columns:
                continue

            # 逐行处理
            for idx, row in df.iterrows():
                # 构建JD文本用于AI分析
                row_text = '\n'.join([f"{col}: {val}" for col, val in row.items() if pd.notna(val)])

                if len(row_text.strip()) < 20:  # 跳过空行或内容太少的行
                    continue

                # 这里返回原始数据，由AI进行解析
                jobs.append({
                    'raw_content': row_text,
                    'source_file': os.path.basename(filepath),
                    'sheet_name': sheet_name,
                    'row_index': idx
                })

        return jobs, None

    except Exception as e:
        return [], f"读取Excel失败: {str(e)}"


def extract_from_text(filepath: str) -> Tuple[List[Dict], Optional[str]]:
    """
    从文本文件提取JD信息
    """
    try:
        with open(filepath, 'r', encoding='utf-8') as f:
            content = f.read()

        if len(content.strip()) < 20:
            return [], "文件内容为空或太短"

        return [{
            'raw_content': content,
            'source_file': os.path.basename(filepath)
        }], None

    except UnicodeDecodeError:
        # 尝试其他编码
        try:
            with open(filepath, 'r', encoding='gbk') as f:
                content = f.read()
            return [{
                'raw_content': content,
                'source_file': os.path.basename(filepath)
            }], None
        except Exception as e:
            return [], f"读取文件失败: {str(e)}"
    except Exception as e:
        return [], f"读取文件失败: {str(e)}"


def extract_from_word(filepath: str) -> Tuple[List[Dict], Optional[str]]:
    """
    从Word文件提取JD信息
    """
    try:
        from docx import Document
        doc = Document(filepath)

        content_parts = []
        for para in doc.paragraphs:
            if para.text.strip():
                content_parts.append(para.text)

        # 也提取表格内容
        for table in doc.tables:
            for row in table.rows:
                row_text = ' | '.join([cell.text.strip() for cell in row.cells if cell.text.strip()])
                if row_text:
                    content_parts.append(row_text)

        content = '\n'.join(content_parts)

        if len(content.strip()) < 20:
            return [], "文件内容为空或太短"

        return [{
            'raw_content': content,
            'source_file': os.path.basename(filepath)
        }], None

    except ImportError:
        return [], "需要安装python-docx库: pip install python-docx"
    except Exception as e:
        return [], f"读取Word文件失败: {str(e)}"


def generate_output_filename(directory: str) -> str:
    """生成输出文件名"""
    date_str = datetime.now().strftime('%Y%m%d')
    return os.path.join(directory, f'jd_summary_{date_str}.xlsx')


def load_existing_data(filepath: str) -> pd.DataFrame:
    """加载已有的Excel数据"""
    if os.path.exists(filepath):
        try:
            return pd.read_excel(filepath)
        except:
            pass
    return pd.DataFrame(columns=OUTPUT_COLUMNS)


def deduplicate_data(existing_df: pd.DataFrame, new_jobs: List[Dict]) -> pd.DataFrame:
    """
    去重合并数据
    以"公司名称 + 职位名称"作为唯一键
    """
    # 将新数据转换为DataFrame
    new_df = pd.DataFrame(new_jobs)

    # 确保所有列都存在
    for col in OUTPUT_COLUMNS:
        if col not in new_df.columns:
            new_df[col] = ''

    new_df = new_df[OUTPUT_COLUMNS]

    if existing_df.empty:
        return new_df

    # 创建唯一键
    def create_key(row):
        company = str(row.get('公司名称', '')).strip()
        position = str(row.get('职位名称', '')).strip()
        return f"{company}_{position}"

    existing_df['_key'] = existing_df.apply(create_key, axis=1)
    new_df['_key'] = new_df.apply(create_key, axis=1)

    # 找出需要更新的和需要新增的
    existing_keys = set(existing_df['_key'].tolist())

    # 更新已有记录
    for idx, row in new_df.iterrows():
        key = row['_key']
        if key in existing_keys:
            # 找到并更新
            mask = existing_df['_key'] == key
            for col in OUTPUT_COLUMNS:
                if col in row and row[col]:
                    existing_df.loc[mask, col] = row[col]

    # 添加新记录
    new_records = new_df[~new_df['_key'].isin(existing_keys)]

    # 合并
    result = pd.concat([existing_df, new_records[OUTPUT_COLUMNS]], ignore_index=True)
    result = result.drop(columns=['_key'], errors='ignore')

    return result


def save_to_excel(df: pd.DataFrame, filepath: str):
    """保存数据到Excel"""
    df.to_excel(filepath, index=False, engine='openpyxl')


def process_directory(directory: str, output_path: Optional[str] = None,
                      append_mode: bool = False) -> Dict:
    """
    处理目录下的所有JD文件

    Args:
        directory: JD文件目录
        output_path: 输出文件路径
        append_mode: 是否追加模式

    Returns:
        处理结果摘要
    """
    # 扫描文件
    files_by_type = scan_directory(directory)

    # 生成输出路径
    if not output_path:
        output_path = generate_output_filename(directory)

    # 收集所有待处理的原始内容
    raw_jobs = []
    skipped_files = []
    file_stats = {
        'excel': 0,
        'word': 0,
        'image': 0,
        'text': 0
    }

    # 处理Excel文件
    for filepath in files_by_type['excel']:
        jobs, error = extract_from_excel(filepath)
        if error:
            skipped_files.append((os.path.basename(filepath), error))
        else:
            raw_jobs.extend(jobs)
            file_stats['excel'] += 1

    # 处理文本文件
    for filepath in files_by_type['text']:
        jobs, error = extract_from_text(filepath)
        if error:
            skipped_files.append((os.path.basename(filepath), error))
        else:
            raw_jobs.extend(jobs)
            file_stats['text'] += 1

    # 处理Word文件
    for filepath in files_by_type['word']:
        jobs, error = extract_from_word(filepath)
        if error:
            skipped_files.append((os.path.basename(filepath), error))
        else:
            raw_jobs.extend(jobs)
            file_stats['word'] += 1

    # 图片文件需要返回路径，由AI视觉模型处理
    image_files = files_by_type['image']
    file_stats['image'] = len(image_files)

    return {
        'raw_jobs': raw_jobs,
        'image_files': image_files,
        'skipped_files': skipped_files,
        'file_stats': file_stats,
        'output_path': output_path,
        'append_mode': append_mode
    }


def create_prompt_for_jd_extraction(content: str, filename: str) -> str:
    """创建用于AI提取JD信息的提示词"""
    return f"""请从以下JD内容中提取信息，按JSON格式返回：

{{
  "岗位分类": "技术/产品/运营/设计/市场/销售/人事/财务/其他",
  "公司名称": "公司名",
  "部门名称": "部门或业务线名称",
  "职位名称": "具体职位名称",
  "薪资范围": "薪资区间或职级",
  "工作地点": "城市",
  "通用硬性要求": "年龄、学历、经验等硬性条件",
  "HR对焦信息描述": "技能要求、目标公司、优先条件等",
  "JD岗位开放日期": "1周内/1月内/2月以上/未知"
}}

JD内容（来源文件：{filename}）：
{content}

注意：
- 如果某个字段信息不存在，返回空字符串""
- 技能关键词、目标公司等信息放在"HR对焦信息描述"中
- 如果有明确的年龄限制、学历要求，放在"通用硬性要求"中
- 只返回JSON，不要有其他内容"""


def create_prompt_for_image_extraction(image_path: str) -> str:
    """创建用于图片分析的提示词"""
    return f"""请分析这张图片（{os.path.basename(image_path)}），如果是JD（职位描述）相关内容（如聊天截图、JD截图等），请提取职位信息。

按以下JSON格式返回：
{{
  "岗位分类": "技术/产品/运营/设计/市场/销售/人事/财务/其他",
  "公司名称": "公司名",
  "部门名称": "部门或业务线名称",
  "职位名称": "具体职位名称",
  "薪资范围": "薪资区间或职级",
  "工作地点": "城市",
  "通用硬性要求": "年龄、学历、经验等硬性条件",
  "HR对焦信息描述": "技能要求、目标公司、优先条件等",
  "JD岗位开放日期": "1周内/1月内/2月以上/未知"
}}

注意：
- 如果图片中有多个职位，请为每个职位单独返回一个JSON对象，放在数组中
- 如果某个字段信息不存在，返回空字符串""
- 如果图片内容与JD无关，返回 {{"error": "非JD内容"}}
- 只返回JSON，不要有其他内容"""


if __name__ == '__main__':
    # 测试代码
    if len(sys.argv) > 1:
        directory = sys.argv[1]
        result = process_directory(directory)
        print(json.dumps({
            'file_stats': result['file_stats'],
            'raw_jobs_count': len(result['raw_jobs']),
            'image_files_count': len(result['image_files']),
            'skipped_files': result['skipped_files'],
            'output_path': result['output_path']
        }, ensure_ascii=False, indent=2))
    else:
        print("Usage: python jd_processor.py <directory>")
